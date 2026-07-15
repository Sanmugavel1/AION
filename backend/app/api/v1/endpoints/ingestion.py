"""
AION API — Document Ingestion
Lets an organization upload real files (policies, finance notes, HR docs,
meeting notes — anything) and turns them into real knowledge the rest of
the platform (OII, disease scan, decay, MRI, board advisor) can read.
"""
from __future__ import annotations

import csv
import io
from typing import Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

from app.ai.llm.llm_client import complete_json
from app.core.dependencies import AuthUser, DbSession
from app.core.logging import get_logger
from app.repositories.graph_repository import GraphRepository
from app.repositories.knowledge_repository import KnowledgeRepository
from app.services.intelligence_index_service import IntelligenceIndexService

logger = get_logger(__name__)

router = APIRouter(prefix="/ingestion", tags=["Document Ingestion"])

MAX_FILE_BYTES = 8 * 1024 * 1024  # 8MB
MAX_EXTRACT_CHARS = 14000  # keep the LLM prompt small and fast
ALLOWED_EXTENSIONS = {"txt", "md", "csv", "pdf", "docx", "xlsx"}

VALID_DOMAINS = {
    "finance", "hr", "policy", "engineering", "product", "sales",
    "operations", "legal", "marketing", "security", "other",
}


class UploadResult(BaseModel):
    knowledge_id: str
    title: str
    domain: str
    plain_summary: str
    relevance_score: float
    tags: list[str]
    source_filename: str
    word_count: int


def _extract_text(filename: str, raw: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '.{ext}'. Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    try:
        if ext in ("txt", "md", "csv"):
            text = raw.decode("utf-8", errors="ignore")
            if ext == "csv":
                # Re-flatten as readable rows rather than a raw comma blob —
                # much easier for the LLM (and a human) to make sense of.
                reader = csv.reader(io.StringIO(text))
                rows = list(reader)
                text = "\n".join(" | ".join(row) for row in rows[:500])
            return text

        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if ext == "docx":
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)

        if ext == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(max_row=500, values_only=True):
                    if any(cell is not None for cell in row):
                        lines.append(" | ".join("" if c is None else str(c) for c in row))
            return "\n".join(lines)

    except Exception as e:
        logger.error("Document text extraction failed", filename=filename, error=str(e))
        raise HTTPException(status_code=422, detail=f"Could not read this file: {e}") from e

    raise HTTPException(status_code=415, detail=f"Unsupported file type '.{ext}'")


EXTRACTION_SYSTEM_PROMPT = (
    "You are Axon, AION's document analyst. An organization has uploaded a real "
    "internal document (it could be about finance, HR policy, salaries, engineering "
    "process, anything). Read it and extract structured knowledge about the "
    "organization. Respond with ONLY a JSON object with these exact keys:\n"
    '  "title": a short, specific title for this document (max 12 words)\n'
    '  "domain": exactly one of: finance, hr, policy, engineering, product, sales, '
    "operations, legal, marketing, security, other\n"
    '  "plain_summary": a 2-4 sentence summary in plain, jargon-free English that any '
    "non-technical employee could understand — explain what this document is and why it matters\n"
    '  "relevance_score": a number from 0.0 to 1.0 estimating how current/actionable this '
    "content is (a stale 3-year-old policy scores low, a current active policy scores high)\n"
    '  "tags": an array of 3-6 short lowercase keyword tags\n'
    "Do not include any text outside the JSON object."
)


@router.post("/upload", response_model=UploadResult)
async def upload_document(
    current_user: AuthUser,
    db: DbSession,
    file: UploadFile = File(...),
):
    """
    Upload any organizational document (policy, finance note, HR record, meeting
    notes, anything). AION reads it, understands it, and turns it into real
    knowledge that the intelligence index, disease scan, decay engine, and board
    advisor all immediately factor in.
    """
    raw = await file.read()
    if len(raw) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File too large — 8MB max")
    if not raw:
        raise HTTPException(status_code=422, detail="File is empty")

    filename = file.filename or "upload"
    text = _extract_text(filename, raw)
    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No readable text found in this file — if it's a scanned/image PDF, "
            "try exporting it as text first.",
        )

    word_count = len(text.split())
    excerpt = text[:MAX_EXTRACT_CHARS]

    try:
        extracted = await complete_json(
            prompt=f"Document filename: {filename}\n\nDocument content:\n{excerpt}",
            system=EXTRACTION_SYSTEM_PROMPT,
        )
    except Exception as e:
        logger.error("Ingestion LLM extraction failed", filename=filename, error=str(e))
        raise HTTPException(
            status_code=502,
            detail="AION couldn't analyze this document right now (AI model unavailable). Please try again shortly.",
        ) from e

    title = str(extracted.get("title") or filename)[:500]
    domain = str(extracted.get("domain") or "other").lower()
    if domain not in VALID_DOMAINS:
        domain = "other"
    plain_summary = str(extracted.get("plain_summary") or "")[:4000]
    try:
        relevance_score = max(0.0, min(1.0, float(extracted.get("relevance_score", 0.7))))
    except (TypeError, ValueError):
        relevance_score = 0.7
    tags = extracted.get("tags") or []
    if not isinstance(tags, list):
        tags = []
    tags = [str(t)[:40] for t in tags[:6]]

    # 1. Persist as a real KnowledgeItem — this is what the OII scorer, decay
    #    engine, and disease detectors all read from.
    k_repo = KnowledgeRepository(db)
    item = await k_repo.create(
        org_id=current_user.org_id,
        creator_id=current_user.user_id,
        title=title,
        content=text[:20000],
        summary=plain_summary,
        domain=domain,
        source_type="document",
        file_type=filename.rsplit(".", 1)[-1].lower() if "." in filename else None,
        relevance_score=relevance_score,
    )
    await db.commit()
    await db.refresh(item)

    # 2. Mirror it into the graph store — this is what MRI, graph search, and
    #    traversal read from. Two stores, kept in sync at write time.
    graph_repo = GraphRepository()
    await graph_repo.create_knowledge_node(
        knowledge_id=str(item.id),
        org_id=current_user.org_id,
        title=title,
        domain=domain,
        source_type="document",
        summary=plain_summary,
        tags=tags,
    )

    logger.info(
        "Document ingested", org_id=current_user.org_id, knowledge_id=str(item.id),
        domain=domain, word_count=word_count,
    )

    # 3. Recompute the Organizational Intelligence Index now, not on the next
    #    scheduled cycle — a customer who just uploaded their first document
    #    expects to immediately see it reflected, not a stale "0% / no data"
    #    reading until some background job gets around to it.
    try:
        from app.api.v1.endpoints.intelligence_index import _compute_live_oii
        intel_service = IntelligenceIndexService(db)
        await _compute_live_oii(intel_service, db, current_user.org_id)
    except Exception as e:
        logger.warning("Post-upload OII refresh failed (non-fatal)", error=str(e))

    return UploadResult(
        knowledge_id=str(item.id),
        title=title,
        domain=domain,
        plain_summary=plain_summary,
        relevance_score=relevance_score,
        tags=tags,
        source_filename=filename,
        word_count=word_count,
    )


class RecentUpload(BaseModel):
    knowledge_id: str
    title: str
    domain: Optional[str]
    plain_summary: Optional[str]
    relevance_score: float
    created_at: str


@router.get("/recent", response_model=list[RecentUpload])
async def list_recent_uploads(current_user: AuthUser, db: DbSession, limit: int = 20):
    """List recently uploaded documents for this organization."""
    k_repo = KnowledgeRepository(db)
    items = await k_repo.get_recent(current_user.org_id, limit=limit)
    return [
        RecentUpload(
            knowledge_id=str(i.id),
            title=i.title,
            domain=i.domain,
            plain_summary=i.summary,
            relevance_score=i.relevance_score,
            created_at=i.created_at.isoformat() if i.created_at else "",
        )
        for i in items
        if i.source_type == "document"
    ]
